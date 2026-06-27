"""Ingestion — arbitrary documents (.md/.txt/.docx/.pdf) → a `Bundle`, plus
normalisation of the loose dict shapes the public API accepts.

Ported from ``src/ingest.py`` (front-matter Markdown, blank-line paragraph
reflow, graceful DOCX/PDF readers). Missing optional deps degrade gracefully.
"""
from __future__ import annotations

import glob
import os
import re

from .models import Bundle, Document, Para

_META_KEYS = {"id", "title", "type", "party", "date"}
_SKIP = {"readme.md", "readme.txt"}

# Filename keyword -> (doc_type, party). Most specific rule first.
_DOC_RULES = [
    ("particulars", "pleading", "claimant"),
    ("claim_form", "pleading", "claimant"),
    ("witness_statement", "witness", "claimant"),
    ("expert_report", "expert", "claimant"),
    ("email", "correspondence", "neutral"),
    ("letter", "correspondence", "neutral"),
    ("defect", "record", "neutral"),
    ("uat", "record", "neutral"),
    ("acceptance", "record", "neutral"),
    ("master_services_agreement", "contract", "neutral"),
    ("statement_of_work", "contract", "neutral"),
    ("order_form", "contract", "neutral"),
    ("deed", "contract", "neutral"),
    ("change_order", "contract", "neutral"),
    ("agreement", "contract", "neutral"),
]


def load_bundle(folder: str) -> Bundle:
    """Load every document file in *folder* into a `Bundle` (sorted by id)."""
    paths: list[str] = []
    for ext in ("*.md", "*.txt", "*.pdf", "*.docx"):
        paths.extend(glob.glob(os.path.join(folder, ext)))
    docs: list[Document] = []
    for p in sorted(paths):
        if os.path.basename(p).lower() in _SKIP:
            continue
        try:
            doc = parse_doc(p)
        except Exception:
            continue
        if doc.paras:
            docs.append(doc)
    docs.sort(key=lambda d: d.id)
    return Bundle(docs=docs)


def parse_doc(path: str) -> Document:
    base = os.path.splitext(os.path.basename(path))[0]
    lower = path.lower()
    if lower.endswith(".docx"):
        return _parse_docx(path, base)
    if lower.endswith(".pdf"):
        raw = _read_pdf(path)
        return _from_text(base, raw)
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        raw = fh.read()
    return _from_text(base, raw)


def _from_text(base: str, raw: str) -> Document:
    meta, body = _split_front_matter(raw)
    doc_type, party = _classify_doc(base)
    return Document(
        id=str(meta.get("id") or _id_from_name(base)),
        title=meta.get("title") or base.replace("_", " "),
        doc_type=meta.get("type") or doc_type,
        party=meta.get("party") or party,
        date=meta.get("date"),
        paras=_paragraphs(body),
    )


def _split_front_matter(raw: str) -> tuple[dict, str]:
    raw = raw.lstrip("﻿").lstrip()
    if not raw.startswith("---"):
        return {}, raw
    parts = raw.split("---", 2)
    if len(parts) < 3:
        return {}, raw
    meta: dict[str, str] = {}
    for line in parts[1].splitlines():
        if ":" in line:
            k, _, v = line.partition(":")
            k = k.strip().lower()
            if k in _META_KEYS:
                meta[k] = v.strip()
    return meta, parts[2]


def _paragraphs(body: str) -> list[Para]:
    chunks = re.split(r"\n\s*\n", body.strip())
    paras: list[Para] = []
    n = 0
    for ch in chunks:
        text = re.sub(r"\s+", " ", ch.strip())
        if not text:
            continue
        n += 1
        paras.append(Para(n=n, text=text))
    return paras


def _classify_doc(base: str) -> tuple[str, str]:
    b = base.lower()
    for kw, doc_type, party in _DOC_RULES:
        if kw in b:
            return doc_type, party
    return "record", "neutral"


def _id_from_name(base: str) -> str:
    m = re.match(r"(\d+)", base)
    return m.group(1) if m else base


def _numbered_paragraphs(lines: list[str]) -> list[Para]:
    paras: list[Para] = []
    cur_n, last, buf = None, 0, []

    def flush() -> None:
        if cur_n is not None:
            text = re.sub(r"\s+", " ", " ".join(buf)).strip()
            if text:
                paras.append(Para(n=cur_n, text=text))

    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        if re.match(r"(?i)^(statement of truth|expert'?s declaration)\b", s):
            break
        m = re.match(r"^(\d{1,3})[.)]\s+(.*)", s)
        if m and int(m.group(1)) > last:
            flush()
            cur_n = last = int(m.group(1))
            buf = [m.group(2)]
        elif cur_n is not None:
            buf.append(s)
    flush()
    return paras


def _seq_paragraphs(lines: list[str]) -> list[Para]:
    out, n = [], 0
    for ln in lines:
        text = re.sub(r"\s+", " ", ln.strip())
        if text:
            n += 1
            out.append(Para(n=n, text=text))
    return out


def _parse_docx(path: str, base: str) -> Document:
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dep
        raise RuntimeError("DOCX support needs `python-docx`.") from exc
    d = docx.Document(path)
    lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            joined = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if joined:
                lines.append(joined)
    doc_type, party = _classify_doc(base)
    paras = _numbered_paragraphs(lines) if doc_type in ("pleading", "witness", "expert") else []
    if not paras:
        paras = _seq_paragraphs(lines)
    title = base.split("_", 1)[-1].replace("_", " ") if "_" in base else base
    return Document(id=_id_from_name(base), title=title, doc_type=doc_type,
                    party=party, date=None, paras=paras)


def _read_pdf(path: str) -> str:  # pragma: no cover - optional dep
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PDF support needs `pdfplumber`.") from exc
    out = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n\n".join(out)


# --------------------------------------------------------------- API coercion
def _coerce_paras(raw) -> list[Para]:
    """Accept paras as ``[(n, text), ...]`` / ``[{n, text}, ...]`` / ``[Para, ...]``."""
    paras: list[Para] = []
    for i, item in enumerate(raw or [], start=1):
        if isinstance(item, Para):
            paras.append(item)
        elif isinstance(item, dict):
            n = int(item.get("n", i))
            paras.append(Para(n=n, text=str(item.get("text", ""))))
        elif isinstance(item, (tuple, list)) and len(item) >= 2:
            paras.append(Para(n=int(item[0]), text=str(item[1])))
        else:
            paras.append(Para(n=i, text=str(item)))
    return paras


def coerce_bundle(bundle) -> Bundle:
    """Normalise the loose ``bundle`` shapes the public API accepts → a `Bundle`.

    Accepts a `Bundle`, or a dict ``{doc_id: {paras, doc_type, party, date, ...}}``
    where ``paras`` entries may be ``(n, text)`` tuples or ``{n, text}`` dicts.
    """
    if isinstance(bundle, Bundle):
        return bundle
    if not isinstance(bundle, dict):
        raise TypeError(f"bundle must be a Bundle or dict, got {type(bundle).__name__}")
    docs: list[Document] = []
    for doc_id, spec in bundle.items():
        spec = spec or {}
        docs.append(Document(
            id=str(doc_id),
            title=str(spec.get("title") or f"Document {doc_id}"),
            doc_type=str(spec.get("doc_type") or "record"),
            party=str(spec.get("party") or "neutral"),
            date=spec.get("date"),
            paras=_coerce_paras(spec.get("paras")),
            category=spec.get("category"),
            modality=str(spec.get("modality") or "document"),
            mime=str(spec.get("mime") or "text/plain"),
            file_url=spec.get("file_url"),
            description=spec.get("description"),
        ))
    docs.sort(key=lambda d: d.id)
    return Bundle(docs=docs)


def coerce_propositions(propositions, bundle: Bundle, *, pleading_tab: str = "02") -> list[dict]:
    """Normalise ``propositions`` → ``[{id, text, pleaded_at:(tab,para)}, ...]``.

    Accepts that exact list-of-dicts shape, or a pleading `Document` whose
    paragraphs each become one proposition pleaded at its own paragraph.
    """
    if isinstance(propositions, Document):
        out = []
        for p in propositions.paras:
            out.append({"id": f"P{p.n}", "text": p.text,
                        "pleaded_at": (propositions.id, p.n)})
        return out
    out = []
    for i, prop in enumerate(propositions or [], start=1):
        if isinstance(prop, dict):
            pid = str(prop.get("id") or f"P{i}")
            text = str(prop.get("text") or "")
            at = prop.get("pleaded_at")
            if isinstance(at, (tuple, list)) and len(at) >= 2:
                pleaded_at = (str(at[0]), int(at[1]))
            else:
                pleaded_at = (pleading_tab, i)
            out.append({"id": pid, "text": text, "pleaded_at": pleaded_at})
        else:  # a Proposition-like object
            out.append({
                "id": str(getattr(prop, "id", f"P{i}")),
                "text": str(getattr(prop, "text", "")),
                "pleaded_at": getattr(prop, "pleaded_at", (pleading_tab, i)),
            })
    return out
