"""Ingestion: a folder of documents -> a `Bundle`.

Now: Markdown / plain-text with a YAML-ish front-matter header. Drop-in for the
real CMS bundle: `.pdf` (via pdfplumber, OCR fallback) and `.docx` (python-docx)
are routed here too; missing optional deps degrade gracefully.

Front-matter format (see data/selftest/bundle/*.md):

    ---
    id: 04
    title: Witness Statement of B. Engineer
    type: witness
    party: defendant
    date: 2018-10-15
    ---
    Paragraph one.

    Paragraph two.

Paragraphs are split on blank lines and numbered 1..N (witness statements are
already numbered; the numbering gives stable "doc ¶n" citation anchors).
"""
from __future__ import annotations

import glob
import hashlib
import json
import os
import re

from .models import Bundle, Document, Para

_META_KEYS = {"id", "title", "type", "party", "date"}
_SKIP = {"readme.md", "readme.txt"}

# Parsing 1.8 GB of PDFs on every run is a non-starter; cache each parse to disk
# keyed on (basename, size, mtime). The cache dir is gitignored.
_CACHE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "cache")


def _cache_file(path: str) -> str:
    st = os.stat(path)
    key = hashlib.md5(f"{os.path.basename(path)}:{st.st_size}:{int(st.st_mtime)}".encode()).hexdigest()
    return os.path.join(_CACHE_DIR, key + ".json")


def _cache_load(path: str) -> Document | None:
    try:
        cf = _cache_file(path)
        if not os.path.exists(cf):
            return None
        with open(cf, "r", encoding="utf-8") as fh:
            d = json.load(fh)
        return Document(id=d["id"], title=d["title"], doc_type=d["doc_type"],
                        party=d["party"], date=d["date"],
                        paras=[Para(p["n"], p["text"]) for p in d["paras"]])
    except Exception:
        return None


def _cache_store(path: str, doc: Document) -> None:
    try:
        os.makedirs(_CACHE_DIR, exist_ok=True)
        with open(_cache_file(path), "w", encoding="utf-8") as fh:
            json.dump({"id": doc.id, "title": doc.title, "doc_type": doc.doc_type,
                       "party": doc.party, "date": doc.date,
                       "paras": [{"n": p.n, "text": p.text} for p in doc.paras]}, fh)
    except Exception:
        pass


def load_bundle(folder: str) -> Bundle:
    """Load every document file in *folder* into a `Bundle` (sorted by id)."""
    paths: list[str] = []
    for ext in ("*.md", "*.txt", "*.pdf", "*.docx"):
        paths.extend(glob.glob(os.path.join(folder, ext)))
    docs: list[Document] = []
    for p in sorted(paths):
        if os.path.basename(p).lower() in _SKIP:
            continue
        doc = parse_doc(p)
        if doc.paras:
            docs.append(doc)
    docs.sort(key=lambda d: d.id)
    return Bundle(docs=docs)


def parse_doc(path: str) -> Document:
    base = os.path.splitext(os.path.basename(path))[0]
    lower = path.lower()
    if lower.endswith(".pdf"):
        # Real CMS bundle: Inquiry witness statements with their own numbered
        # paragraphs. Parse those as citation anchors (not a blank-line reflow).
        cached = _cache_load(path)
        if cached is not None:
            return cached
        doc = parse_statement(_read_pdf(path), fallback_id=base)
        _cache_store(path, doc)
        return doc
    if lower.endswith(".docx"):
        return _parse_bundle_docx(path, base)
    else:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            raw = fh.read()
        meta, raw = _split_front_matter(raw)

    paras = _paragraphs(raw)
    doc_id = meta.get("id") or _id_from_name(base)
    return Document(
        id=str(doc_id),
        title=meta.get("title") or base.replace("_", " "),
        doc_type=meta.get("type") or "unknown",
        party=meta.get("party") or "neutral",
        date=meta.get("date"),
        paras=paras,
    )


# --------------------------------------------------- witness-statement parser
# Statement numbers are uniformly "WITN" + 8 digits.
_WITN = re.compile(r"WITN\d{8}")
_PAGE = re.compile(r"(?i)^page\s+\d+\s+of\s*\d+\s*$")
_PARA = re.compile(r"^(\d{1,3})[.)]\s+(.*)")
_HEADING = re.compile(r"^[A-Z][A-Z0-9 &'/().,\-]{1,48}$")


def _clean_id(name: str) -> str:
    """Stable doc id from a filename/header: the 'WITN'+8-digit statement number."""
    t = re.sub(r"[^A-Za-z0-9]", "", name).upper().replace("O", "0")
    m = _WITN.search(t)
    return m.group(0) if m else (name.strip() or "DOC")


def _is_watermark(s: str) -> bool:
    """A line that is only the (often OCR-mangled) statement-number watermark."""
    t = re.sub(r"[^A-Za-z0-9]", "", s).upper().replace("O", "0")
    return bool(t) and bool(re.fullmatch(r"WITN\d{6,}", t))


def _is_heading(s: str) -> bool:
    """ALL-CAPS section heading between paragraphs (INTRODUCTION, BACKGROUND)."""
    if not _HEADING.fullmatch(s):
        return False
    return len(s.split()) <= 7 and not s.endswith(".")


def parse_statement(raw: str, *, fallback_id: str = "") -> Document:
    """Parse one Inquiry witness statement into a `Document` with REAL paragraph
    numbers as anchors. Pure function over pdfplumber's extracted text.

    - metadata from the header (Witness Name / Statement No. / Dated / "... STATEMENT OF X");
    - drops watermark + "Page X of Y" lines and ALL-CAPS section headings;
    - joins multi-line paragraphs, including across page breaks;
    - stops the body at "Statement of Truth".
    Falls back to a blank-line reflow if the document has no numbered paragraphs.
    """
    lines = raw.splitlines()
    meta: dict[str, str] = {}
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        m = re.match(r"(?i)^witness name:?\s*(.+)", s)
        if m and "name" not in meta:
            meta["name"] = m.group(1).strip()
        m = re.match(r"(?i)^statement no\.?:?\s*(.+)", s)
        if m and "stmt" not in meta:
            meta["stmt"] = _clean_id(m.group(1))
        m = re.match(r"(?i)^dated\s*:?\s*(.+)", s)
        if m and "date" not in meta:
            meta["date"] = re.sub(r"^[:\s]+", "", m.group(1)).strip()
        m = re.match(r"(?i)^(?:first|second|third|fourth|fifth|further)?\s*witness "
                     r"statement of\s+(.+)", s)
        if m and "name" not in meta:
            meta["name"] = m.group(1).strip().title()

    paras: list[Para] = []
    cur_n: int | None = None
    buf: list[str] = []

    def flush() -> None:
        if cur_n is not None:
            text = re.sub(r"\s+", " ", " ".join(buf)).strip()
            if text:
                paras.append(Para(n=cur_n, text=text))

    for ln in lines:
        s = ln.strip()
        if not s or _PAGE.match(s) or _is_watermark(s):
            continue
        if re.match(r"(?i)^statement of truth", s):
            break
        m = _PARA.match(s)
        if m:
            flush()
            cur_n, buf = int(m.group(1)), [m.group(2)]
        elif cur_n is not None and not _is_heading(s):
            buf.append(s)
    flush()

    doc_id = meta.get("stmt") or _clean_id(fallback_id)
    title = meta.get("name") or _clean_id(fallback_id)
    if not paras:  # no numbered paragraphs — fall back to blank-line reflow
        paras = _paragraphs(raw)
    return Document(id=doc_id, title=title, doc_type="witness",
                    party="neutral", date=meta.get("date"), paras=paras)


def _split_front_matter(raw: str) -> tuple[dict, str]:
    raw = raw.lstrip("﻿").lstrip()
    if not raw.startswith("---"):
        return {}, raw
    parts = raw.split("---", 2)
    if len(parts) < 3:
        return {}, raw
    meta: dict[str, str] = {}
    for line in parts[1].splitlines():
        if ":" in line:
            k, _, v = line.partition(":")
            k = k.strip().lower()
            if k in _META_KEYS:
                meta[k] = v.strip()
    return meta, parts[2]


def _paragraphs(body: str) -> list[Para]:
    chunks = re.split(r"\n\s*\n", body.strip())
    paras: list[Para] = []
    n = 0
    for ch in chunks:
        text = re.sub(r"\s+", " ", ch.strip())
        if not text:
            continue
        n += 1
        paras.append(Para(n=n, text=text))
    return paras


def _id_from_name(base: str) -> str:
    m = re.match(r"(\d+)", base)
    return m.group(1) if m else base


# ----------------------------------------------------- DOCX litigation bundle
# Filename -> (doc_type, party). Most specific rule wins, so side-specific
# correspondence is listed before the generic "letter"/"email" fallback.
_DOC_RULES = [
    ("particulars", "pleading", "claimant"),
    ("claim_form", "pleading", "claimant"),
    ("witness_statement", "witness", "claimant"),
    ("expert_report", "expert", "claimant"),
    ("techflow_response", "correspondence", "defendant"),
    ("notice_of_termination", "correspondence", "claimant"),
    ("email", "correspondence", "neutral"),
    ("letter", "correspondence", "neutral"),
    ("defect", "record", "neutral"),
    ("uat", "record", "neutral"),
    ("acceptance", "record", "neutral"),
    ("master_services_agreement", "contract", "neutral"),
    ("statement_of_work", "contract", "neutral"),
    ("order_form", "contract", "neutral"),
    ("deed", "contract", "neutral"),
    ("change_order", "contract", "neutral"),
    ("agreement", "contract", "neutral"),
    ("bundle_index", "index", "neutral"),
]


def _classify_doc(base: str) -> tuple[str, str]:
    b = base.lower()
    for kw, doc_type, party in _DOC_RULES:
        if kw in b:
            return doc_type, party
    return "unknown", "neutral"


def _numbered_paragraphs(lines: list[str]) -> list[Para]:
    """Group lines into paragraphs by the document's OWN numbering ("1.", "2." ...).

    A numbered line opens a new paragraph only when its number is strictly
    greater than the last opened one; this keeps witness/pleading numbering exact
    (consecutive or gapped) and lets expert-report section headings (which restart
    at 1) fold into the running paragraph instead of derailing it. Body stops at
    the truth/declaration block.
    """
    paras: list[Para] = []
    cur_n: int | None = None
    last = 0
    buf: list[str] = []

    def flush() -> None:
        if cur_n is not None:
            text = re.sub(r"\s+", " ", " ".join(buf)).strip()
            if text:
                paras.append(Para(n=cur_n, text=text))

    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        if re.match(r"(?i)^(statement of truth|expert'?s declaration)\b", s):
            break
        m = re.match(r"^(\d{1,3})[.)]\s+(.*)", s)
        if m and int(m.group(1)) > last:
            flush()
            cur_n = last = int(m.group(1))
            buf = [m.group(2)]
        elif cur_n is not None:
            buf.append(s)
    flush()
    return paras


def _seq_paragraphs(lines: list[str]) -> list[Para]:
    out: list[Para] = []
    n = 0
    for ln in lines:
        text = re.sub(r"\s+", " ", ln.strip())
        if text:
            n += 1
            out.append(Para(n=n, text=text))
    return out


def _docx_lines(path: str) -> list[str]:
    """Body paragraphs + table rows (tables hold key evidence: the defect log,
    email headers). Each non-empty paragraph and each table row is one line."""
    try:
        import docx  # type: ignore
    except ImportError:
        raise RuntimeError("DOCX support needs `python-docx` (pip install python-docx).")
    d = docx.Document(path)
    lines: list[str] = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            joined = " | ".join(c for c in cells if c)
            if joined:
                lines.append(joined)
    return lines


def _parse_bundle_docx(path: str, base: str) -> Document:
    lines = _docx_lines(path)
    doc_type, party = _classify_doc(base)
    paras: list[Para] = []
    if doc_type in ("pleading", "witness", "expert"):
        paras = _numbered_paragraphs(lines)
    if not paras:                      # contracts/records/correspondence, or no numbering
        paras = _seq_paragraphs(lines)
    title = base.split("_", 1)[-1].replace("_", " ") if "_" in base else base
    return Document(id=_id_from_name(base), title=title, doc_type=doc_type,
                    party=party, date=None, paras=paras)


# ------------------------------------------------------------- format readers
def _read_pdf(path: str) -> str:
    try:
        import pdfplumber  # type: ignore
        out = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                if not txt.strip():
                    txt = _ocr_page(page)
                out.append(txt)
        return "\n\n".join(out)
    except ImportError:
        raise RuntimeError(
            "PDF support needs `pdfplumber` (pip install pdfplumber); for scanned "
            "PDFs also `pytesseract` + the Tesseract binary."
        )


def _ocr_page(page) -> str:
    """Best-effort OCR of a scanned page; silently returns '' if OCR unavailable."""
    try:
        import pytesseract  # type: ignore
        img = page.to_image(resolution=200).original
        return pytesseract.image_to_string(img) or ""
    except Exception:
        return ""


def _read_docx(path: str) -> str:
    try:
        import docx  # type: ignore
    except ImportError:
        raise RuntimeError("DOCX support needs `python-docx` (pip install python-docx).")
    d = docx.Document(path)
    return "\n\n".join(p.text for p in d.paragraphs if p.text.strip())
